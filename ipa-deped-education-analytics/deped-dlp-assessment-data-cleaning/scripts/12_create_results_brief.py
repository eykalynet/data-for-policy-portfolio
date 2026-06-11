################################################################################
# TITLE   : 12_create_results_brief.py
# PURPOSE : Make a short Word brief that explains the main descriptive results.
# PROJECT : Dynamic Learning Program descriptive results
# AUTHOR  : Erika Salvador
# DATE    : June 11, 2026
################################################################################

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path.cwd()
if PROJECT_DIR.name == "scripts":
    PROJECT_DIR = PROJECT_DIR.parent

TABLE_DIR = PROJECT_DIR / "outputs" / "tables"
FIGURE_DIR = PROJECT_DIR / "outputs" / "figures" / "individual_figures"
OUTPUT_DOCX = PROJECT_DIR / "outputs" / "descriptive_results_brief.docx"


def fmt_int(x):
    return f"{int(round(float(x))):,}"


def fmt_pct(x):
    return f"{float(x):.1f}%"


def read_excel(name):
    return pd.read_excel(TABLE_DIR / name)


def read_csv(name):
    return pd.read_csv(TABLE_DIR / name)


def set_doc_style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Georgia"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E5E9E"),
        ("Heading 2", 13, "2E5E9E"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DLP Assessment Descriptive Results")
    run.font.name = "Georgia"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E5E9E")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("School-level snapshot, compliance, proficiency, and geographic summaries")
    run.font.name = "Georgia"
    run.font.size = Pt(11)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("June 11, 2026")
    run.font.name = "Georgia"
    run.font.size = Pt(10)


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Georgia"
                run.font.bold = True
                run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc, title, filename, width=6.6):
    path = FIGURE_DIR / filename
    if not path.exists():
        return
    caption = doc.add_paragraph()
    caption.style = doc.styles["Heading 3"]
    caption.add_run(title)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def rma_not_proficient_rows(data):
    table = data[data["proficiency_group"] == "not proficient"].copy()
    table = (
        table.groupby(["treatment_status", "time_point"], as_index=False)
        .agg(student_count=("student_count", "sum"), assessed_count=("assessed_count", "sum"))
    )
    table["percent"] = 100 * table["student_count"] / table["assessed_count"]
    wide = table.pivot(index="treatment_status", columns="time_point", values="percent").reset_index()
    rows = []
    for _, row in wide.iterrows():
        rows.append([
            row["treatment_status"],
            fmt_pct(row.get("BoSY", 0)),
            fmt_pct(row.get("EoSY", 0)),
        ])
    return rows


def philiri_grade_ready_rows(data):
    table = data[data["reading_category"] == "independent"].copy()
    table = table[
        ((table["time_point"] == "BoSY") & (table["level_group"].isin(["2-level", "3-level"])))
        | (table["time_point"] == "EoSY")
    ].copy()
    table["reading_measure"] = ""
    table.loc[(table["time_point"] == "BoSY") & (table["level_group"] == "2-level"), "reading_measure"] = "BoSY 2-Level"
    table.loc[(table["time_point"] == "BoSY") & (table["level_group"] == "3-level"), "reading_measure"] = "BoSY 3-Level"
    table.loc[table["time_point"] == "EoSY", "reading_measure"] = "EoSY Proxy"
    table = (
        table.groupby(["treatment_status", "reading_measure"], as_index=False)
        .agg(student_count=("student_count", "sum"), total_classified=("total_classified", "sum"))
    )
    table["percent"] = 100 * table["student_count"] / table["total_classified"]
    wide = table.pivot(index="treatment_status", columns="reading_measure", values="percent").reset_index()
    rows = []
    for _, row in wide.iterrows():
        rows.append([
            row["treatment_status"],
            fmt_pct(row.get("BoSY 2-Level", 0)),
            fmt_pct(row.get("BoSY 3-Level", 0)),
            fmt_pct(row.get("EoSY Proxy", 0)),
        ])
    return rows


def main():
    snapshot = read_excel("04_overall_dataset_snapshot.xlsx").iloc[0]
    treatment = read_excel("04_treatment_status_summary.xlsx")
    enrollment = read_excel("04_enrollment_assessment_summary.xlsx")
    assignment = read_excel("05_compliance_summary_assignment_group.xlsx")
    region_compliance = read_excel("05_compliance_summary_region.xlsx")
    geography = read_csv("07_geographic_summary_region_map_ready.csv")
    rma = read_excel("04_rma_proficiency_by_grade.xlsx")
    philiri = read_excel("04_philiri_reading_by_grade.xlsx")
    rma_treatment = read_excel("04_rma_proficiency_by_treatment_status.xlsx")
    philiri_treatment = read_excel("04_philiri_reading_by_treatment_status.xlsx")

    total_schools = int(snapshot["n_schools"])
    control_schools = int(treatment.loc[treatment["treatment_status"] == "Control", "n_schools"].iloc[0])
    treatment_schools = int(treatment.loc[treatment["treatment_status"] == "Treatment", "n_schools"].iloc[0])
    compliant = int(assignment["n_compliant"].sum())
    compliance_rate = 100 * compliant / total_schools

    rma_total = (
        rma.groupby(["time_point", "proficiency_group"], as_index=False)
        .agg(student_count=("student_count", "sum"), assessed_count=("assessed_count", "sum"))
    )
    rma_total["percent"] = 100 * rma_total["student_count"] / rma_total["assessed_count"]
    rma_not = rma_total[rma_total["proficiency_group"] == "not proficient"]

    philiri_total = (
        philiri.groupby(["time_point", "level_group", "reading_category"], as_index=False)
        .agg(student_count=("student_count", "sum"), total_classified=("total_classified", "sum"))
    )
    philiri_total["percent"] = 100 * philiri_total["student_count"] / philiri_total["total_classified"]

    doc = Document()
    set_doc_style(doc)
    add_title(doc)

    doc.add_heading("What is in the dataset", level=1)
    doc.add_paragraph(
        f"The final file uses the DLP randomized schools file as the base, so the dataset has "
        f"{fmt_int(total_schools)} schools. Based on rev_status, {fmt_int(control_schools)} schools "
        f"are coded as Control and {fmt_int(treatment_schools)} schools are coded as Treatment."
    )
    doc.add_paragraph(
        f"Across these schools, the DLP enrollment file has {fmt_int(snapshot['enrolled_jhs'])} "
        f"junior high school students. The Beginning of School Year reading assessment has "
        f"{fmt_int(snapshot['philiri_bosy_assessed'])} assessed students, while the End of School Year "
        f"reading assessment has {fmt_int(snapshot['philiri_eosy_assessed'])}. For mathematics, "
        f"the Beginning of School Year assessment has {fmt_int(snapshot['rma_bosy_assessed'])} assessed "
        f"students and the End of School Year assessment has {fmt_int(snapshot['rma_eosy_assessed'])}."
    )
    add_table(
        doc,
        [
            [
                row["treatment_status"],
                fmt_int(row["n_schools"]),
                fmt_int(row["enrolled_jhs"]),
                fmt_int(row["philiri_bosy_assessed"]),
                fmt_int(row["philiri_eosy_assessed"]),
                fmt_int(row["rma_bosy_assessed"]),
                fmt_int(row["rma_eosy_assessed"]),
            ]
            for _, row in enrollment.iterrows()
        ],
        [
            "Treatment status",
            "Schools",
            "Enrollment",
            "Reading BoSY assessed",
            "Reading EoSY assessed",
            "Math BoSY assessed",
            "Math EoSY assessed",
        ],
    )
    add_figure(doc, "Dataset snapshot by treatment status", "06_dataset_snapshot_treatment_status.png")

    doc.add_heading("Compliance", level=1)
    doc.add_paragraph(
        f"Using the assignment flags and rev_status rule, {fmt_int(compliant)} of {fmt_int(total_schools)} "
        f"schools are compliant, or {fmt_pct(compliance_rate)} overall. Compliance is highest in the "
        f"emergency assignment group and lowest in the mainstream assignment group."
    )
    add_table(
        doc,
        [
            [
                row["assignment_group"].title(),
                fmt_int(row["n_schools"]),
                fmt_int(row["n_compliant"]),
                fmt_pct(row["compliance_rate"]),
            ]
            for _, row in assignment.iterrows()
        ],
        ["Assignment group", "Schools", "Compliant schools", "Compliance rate"],
    )
    add_figure(doc, "Compliance by assigned randomization group", "06_compliance_by_assignment_group.png")

    doc.add_heading("Proficiency and reading categories", level=1)
    doc.add_paragraph(
        "Rapid Mathematics Assessment results are concentrated in the not proficient group in both "
        "the Beginning of School Year and End of School Year files. This pattern is visible across "
        "Grades 7 to 10."
    )
    add_table(
        doc,
        [
            [
                row["time_point"],
                "Not proficient",
                fmt_int(row["student_count"]),
                fmt_pct(row["percent"]),
            ]
            for _, row in rma_not.iterrows()
        ],
        ["Period", "Group", "Students", "Share of assessed"],
    )
    doc.add_paragraph(
        "For the Philippine Informal Reading Inventory, the Beginning of School Year file has separate "
        "2-level and 3-level reading categories. The End of School Year file does not have the same "
        "2-level and 3-level fields, so End of School Year independent is used as a proxy for grade ready "
        "in the comparison figures."
    )
    bosy_two = philiri_total[
        (philiri_total["time_point"] == "BoSY")
        & (philiri_total["level_group"] == "2-level")
        & (philiri_total["reading_category"] == "independent")
    ].iloc[0]
    bosy_three = philiri_total[
        (philiri_total["time_point"] == "BoSY")
        & (philiri_total["level_group"] == "3-level")
        & (philiri_total["reading_category"] == "independent")
    ].iloc[0]
    eosy_independent = philiri_total[
        (philiri_total["time_point"] == "EoSY")
        & (philiri_total["reading_category"] == "independent")
    ].iloc[0]
    add_table(
        doc,
        [
            ["Beginning of School Year 2-Level", "Independent", fmt_pct(bosy_two["percent"])],
            ["Beginning of School Year 3-Level", "Independent", fmt_pct(bosy_three["percent"])],
            ["End of School Year Proxy", "Independent", fmt_pct(eosy_independent["percent"])],
        ],
        ["Reading measure", "Category used", "Share of classified students"],
    )
    add_figure(doc, "Rapid Mathematics Assessment proficiency by grade", "06_rma_proficiency_bosy_eosy_by_grade.png")
    add_figure(
        doc,
        "Philippine Informal Reading Inventory Beginning of School Year 2-Level and End of School Year Proxy",
        "06_philiri_reading_categories_bosy_2level_eosy_proxy_by_grade.png",
    )
    add_figure(
        doc,
        "Philippine Informal Reading Inventory Beginning of School Year 3-Level and End of School Year Proxy",
        "06_philiri_reading_categories_bosy_3level_eosy_proxy_by_grade.png",
    )
    doc.add_heading("Subgroup patterns", level=1)
    doc.add_paragraph(
        "The subgroup tables below compare the main math and reading indicators for Control and Treatment schools. "
        "The full grade-level subgroup tables are saved in the outputs/tables folder."
    )
    doc.add_heading("RMA not proficient share", level=2)
    add_table(
        doc,
        rma_not_proficient_rows(rma_treatment),
        ["Treatment status", "Beginning of School Year", "End of School Year"],
    )
    add_figure(
        doc,
        "RMA not proficient share by treatment status",
        "06_subgroup_rma_not_proficient_by_treatment_status.png",
    )

    doc.add_heading("Phil-IRI grade ready and proxy share", level=2)
    doc.add_paragraph(
        "For Phil-IRI, Beginning of School Year uses the 2-level and 3-level independent categories. "
        "End of School Year independent is shown as a proxy because the End of School Year file does not "
        "include the same 2-level and 3-level fields."
    )
    add_table(
        doc,
        philiri_grade_ready_rows(philiri_treatment),
        ["Treatment status", "BoSY 2-Level", "BoSY 3-Level", "EoSY Proxy"],
    )
    add_figure(
        doc,
        "Phil-IRI grade ready share by treatment status",
        "06_subgroup_philiri_grade_ready_by_treatment_status.png",
    )

    doc.add_heading("Geographic summaries", level=1)
    doc.add_paragraph(
        "The randomized schools in this file are spread across five regions. Region 5 has the largest "
        "number of schools in the dataset. The coverage rates now use the enrollment denominators from "
        "the assessment dashboard files, so the regional rates stay below 100 percent."
    )
    geography = geography.sort_values("n_schools", ascending=False)
    add_table(
        doc,
        [
            [
                f"Region {int(row['region_code'])}",
                fmt_int(row["n_schools"]),
                fmt_pct(row["philiri_bosy_assessment_rate"]),
                fmt_pct(row["philiri_eosy_assessment_rate"]),
                fmt_pct(row["rma_bosy_assessment_rate"]),
                fmt_pct(row["rma_eosy_assessment_rate"]),
                fmt_pct(row["compliance_rate"]),
            ]
            for _, row in geography.iterrows()
        ],
        [
            "Region",
            "Schools",
            "Reading BoSY coverage",
            "Reading EoSY coverage",
            "Math BoSY coverage",
            "Math EoSY coverage",
            "Compliance",
        ],
    )
    add_figure(doc, "Beginning and End of School Year regional assessment coverage", "07_regional_assessment_coverage_bosy_eosy.png")
    add_figure(doc, "Regional compliance rate", "07_regional_compliance_rate.png")

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.add_heading("Files to share", level=1)
    doc.add_paragraph(
        "The do-files are in scripts/. Cleaned data is in data/. Tables, logs, this brief, and individual "
        "PNG figures are in outputs/."
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
